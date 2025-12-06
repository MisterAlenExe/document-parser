"""PyMuPDF4LLM-based document extractor with Tesseract OCR support."""

import io
import os
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pymupdf4llm
import pytesseract
from docx import Document
from PIL import Image

from document_parser.config import OCRConfig
from document_parser.exceptions import ExtractionError
from document_parser.logger import Timer, get_logger

logger = get_logger(__name__)


class PyMuPDFExtractor:
    """Lightweight document extractor using PyMuPDF4LLM.

    No AI/ML dependencies - uses PyMuPDF for PDFs, python-docx for DOCX,
    and Tesseract OCR for images. Outputs markdown format optimized for LLMs.
    """

    def __init__(self, config: Optional[OCRConfig] = None):
        """Initialize extractor with configuration.

        Args:
            config: OCR configuration. If None, uses defaults.
        """
        self.config = config or OCRConfig()
        self.last_ocr_used = False

        # Set Tesseract environment variables if provided
        if self.config.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = self.config.tesseract_cmd
        if self.config.tessdata_prefix:
            os.environ["TESSDATA_PREFIX"] = self.config.tessdata_prefix

        logger.info(
            "Initializing PyMuPDFExtractor",
            extra_data={
                "tesseract_cmd": self.config.tesseract_cmd,
                "languages": self.config.languages,
                "dpi": self.config.dpi,
            },
        )

    def extract(self, file_bytes: bytes, file_name: str) -> str:
        """Extract text from document.

        Args:
            file_bytes: Raw file bytes
            file_name: Original filename (used for format detection)

        Returns:
            Extracted text in Markdown format

        Raises:
            ExtractionError: If extraction fails
        """
        suffix = Path(file_name).suffix.lower()
        self.last_ocr_used = False

        logger.debug(
            "Starting document extraction",
            extra_data={
                "file_name": file_name,
                "file_extension": suffix,
                "file_size_bytes": len(file_bytes),
            },
        )

        try:
            if suffix == ".pdf":
                return self._extract_pdf(file_bytes, file_name)
            elif suffix == ".docx":
                return self._extract_docx(file_bytes, file_name)
            elif suffix == ".doc":
                return self._extract_doc(file_bytes, file_name)
            elif suffix == ".xls":
                return self._extract_xls(file_bytes, file_name)
            elif suffix == ".xlsx":
                return self._extract_xlsx(file_bytes, file_name)
            elif suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                return self._extract_image(file_bytes, file_name)
            else:
                # Fallback: treat as PDF
                logger.warning(
                    "Unknown file extension, attempting PDF extraction",
                    extra_data={
                        "file_name": file_name,
                        "file_extension": suffix,
                    },
                )
                return self._extract_pdf(file_bytes, file_name)

        except Exception as exc:
            logger.error(
                "Document extraction failed",
                extra_data={
                    "file_name": file_name,
                    "file_extension": suffix,
                    "error_type": type(exc).__name__,
                    "error": str(exc),
                },
                exc_info=True,
            )
            raise ExtractionError(f"Failed to extract document: {exc}") from exc

    def _extract_pdf(self, file_bytes: bytes, file_name: str = "unknown.pdf") -> str:
        """Extract from PDF using PyMuPDF4LLM."""
        # pymupdf4llm requires a file path, so we still need temp file for native extraction
        # But we can optimize OCR path to potentially use in-memory handling
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            page_count = self._get_page_count(tmp_path)

            # PyMuPDF4LLM handles:
            # - Native text extraction (fast path)
            # - Table detection and markdown formatting
            # - Structured output (headers, lists, etc.)
            # Note: Does NOT include OCR - scanned PDFs will return empty text
            with Timer("pdf_native_extraction") as native_timer:
                md_text = pymupdf4llm.to_markdown(
                    tmp_path,
                    # Table extraction
                    table_strategy="lines_strict",  # Detect tables with visible lines
                    # Text extraction
                    force_text=True,  # Extract text even over images
                    # Image handling
                    write_images=False,  # Don't save images separately
                    ignore_images=True,  # Skip decorative images
                    # Text processing
                    ignore_code=False,  # Preserve code formatting
                    fontsize_limit=3,  # Ignore text smaller than 3pt
                )

            text = md_text.strip()
            char_count = len(text)

            logger.debug(
                "PDF native text extraction completed",
                extra_data={
                    "file_name": file_name,
                    "characters_extracted": char_count,
                    "page_count": page_count,
                    "extraction_time_ms": native_timer.get_elapsed_ms(),
                },
            )

            should_ocr = self._should_ocr_pdf(char_count, page_count, len(file_bytes))

            # If no text extracted or text clearly too small for PDF size/page count
            if should_ocr:
                logger.info(
                    "Triggering OCR fallback for PDF",
                    extra_data={
                        "file_name": file_name,
                        "native_characters": char_count,
                        "page_count": page_count,
                    },
                )

                with Timer("pdf_ocr") as ocr_timer:
                    # Use existing temp file path for OCR (already on disk)
                    ocr_text = self._ocr_pdf(tmp_path, file_name)

                logger.info(
                    "OCR extraction completed",
                    extra_data={
                        "file_name": file_name,
                        "characters_extracted": len(ocr_text),
                        "ocr_time_ms": ocr_timer.get_elapsed_ms(),
                    },
                )

                # Prefer OCR output if it is longer than the native extraction
                if len(ocr_text) > char_count:
                    text = ocr_text
                    self.last_ocr_used = True

            return text

        finally:
            # Cleanup temp file
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass

    def _ocr_page(
        self, pdf_path: str, page_num: int, file_name: str
    ) -> tuple[int, str, float]:
        """OCR a single page - worker function for parallel processing.

        Args:
            pdf_path: Path to the PDF file
            page_num: Page number to process (0-indexed)
            file_name: Original file name for logging

        Returns:
            Tuple of (page_num, extracted_text, elapsed_time_ms)
        """
        import time

        start_time = time.time()

        try:
            # Open PDF and render specific page
            pdf_document = fitz.open(pdf_path)
            page = pdf_document[page_num]
            pix = page.get_pixmap(dpi=self.config.dpi)

            # Convert to PIL Image
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))

            # Perform OCR with configured languages
            page_text = pytesseract.image_to_string(
                image,
                lang=self.config.languages,
                config=f"--psm {self.config.psm_mode}",
            )

            pdf_document.close()

            elapsed_ms = (time.time() - start_time) * 1000

            logger.info(
                f"OCR completed for page {page_num + 1}",
                extra_data={
                    "file_name": file_name,
                    "page_number": page_num + 1,
                    "characters_extracted": len(page_text.strip()),
                    "ocr_time_ms": elapsed_ms,
                },
            )

            return (page_num, page_text.strip(), elapsed_ms)

        except Exception as e:
            logger.error(
                f"OCR failed for page {page_num + 1}",
                extra_data={
                    "file_name": file_name,
                    "page_number": page_num + 1,
                    "error_type": type(e).__name__,
                    "error": str(e),
                },
                exc_info=True,
            )
            return (page_num, "", 0.0)

    def _ocr_pdf(self, pdf_path: str, file_name: str = "unknown.pdf") -> str:
        """Perform OCR on a PDF using Tesseract with parallel processing.

        Args:
            pdf_path: Path to the PDF file
            file_name: Original file name for logging

        Returns:
            Extracted text from all pages
        """
        try:
            # Get page count
            pdf_document = fitz.open(pdf_path)
            page_count = len(pdf_document)
            pdf_document.close()

            logger.debug(
                "Starting parallel OCR on PDF pages",
                extra_data={
                    "file_name": file_name,
                    "page_count": page_count,
                    "max_workers": self.config.max_workers,
                },
            )

            # Process pages in parallel using ThreadPoolExecutor
            page_results = {}
            with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
                # Submit all pages for processing
                future_to_page = {
                    executor.submit(self._ocr_page, pdf_path, page_num, file_name): page_num
                    for page_num in range(page_count)
                }

                # Collect results as they complete
                for future in as_completed(future_to_page):
                    page_num, page_text, _ = future.result()
                    page_results[page_num] = page_text

            # Reconstruct text in correct page order
            all_text = [page_results[i] for i in range(page_count) if page_results[i]]

            total_text = "\n\n".join(all_text)

            logger.info(
                "PDF OCR completed for all pages",
                extra_data={
                    "file_name": file_name,
                    "page_count": page_count,
                    "pages_with_text": len(all_text),
                    "total_characters": len(total_text),
                },
            )

            return total_text

        except Exception as e:
            logger.error(
                "OCR failed for PDF",
                extra_data={
                    "file_name": file_name,
                    "error_type": type(e).__name__,
                    "error": str(e),
                },
                exc_info=True,
            )
            return ""

    def _should_ocr_pdf(
        self, native_char_count: int, page_count: int, file_size_bytes: int
    ) -> bool:
        """Decide whether to run OCR fallback after native extraction."""

        # If no text was extracted, always try OCR
        if native_char_count == 0:
            return True

        # Heuristic: very little text per page likely means scanned PDF
        if page_count > 0 and (
            native_char_count / page_count
        ) < self.config.pdf_ocr_min_chars_per_page:
            return True

        # Heuristic: small absolute text for a reasonably sized file
        if (
            native_char_count < self.config.pdf_ocr_min_chars
            and file_size_bytes >= self.config.pdf_ocr_min_file_size_bytes
        ):
            return True

        return False

    @staticmethod
    def _get_page_count(pdf_path: str) -> int:
        """Return page count, swallowing errors to keep extraction resilient."""
        try:
            pdf_document = fitz.open(pdf_path)
            page_count = len(pdf_document)
            pdf_document.close()
            return page_count
        except Exception:
            return 0

    def _extract_docx(self, file_bytes: bytes, file_name: str = "unknown.docx") -> str:
        """Extract from DOCX using python-docx."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            with Timer("docx_extraction") as timer:
                doc = Document(tmp_path)

                # Extract paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

                # Extract tables (markdown format)
                tables = []
                for table in doc.tables:
                    # Build markdown table
                    rows = []
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))

                        # Add header separator after first row
                        if i == 0:
                            separator = " | ".join(["---"] * len(cells))
                            rows.append(separator)

                    if rows:
                        tables.append("\n".join(rows))

                # Combine paragraphs and tables
                parts = paragraphs
                if tables:
                    parts.extend(tables)

                result = "\n\n".join(parts)

            logger.debug(
                "DOCX extraction completed",
                extra_data={
                    "file_name": file_name,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables),
                    "characters_extracted": len(result),
                    "extraction_time_ms": timer.get_elapsed_ms(),
                },
            )

            return result

        finally:
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass

    def _extract_doc(self, file_bytes: bytes, file_name: str = "unknown.doc") -> str:
        """Extract text from legacy .doc using system converters if available."""
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = Path(tmp_file.name)

        try:
            # Prefer macOS textutil if present
            if shutil.which("textutil"):
                with Timer("doc_textutil") as timer:
                    result = subprocess.run(
                        [
                            "textutil",
                            "-convert",
                            "txt",
                            str(tmp_path),
                            "-stdout",
                        ],
                        capture_output=True,
                        text=True,
                    )

                if result.returncode == 0 and result.stdout.strip():
                    logger.info(
                        "DOC extraction completed via textutil",
                        extra_data={
                            "file_name": file_name,
                            "characters_extracted": len(result.stdout.strip()),
                            "extraction_time_ms": timer.get_elapsed_ms(),
                        },
                    )
                    return result.stdout.strip()

            # Fallback to LibreOffice/soffice if available
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if soffice:
                with tempfile.TemporaryDirectory() as tmp_dir:
                    with Timer("doc_soffice") as timer:
                        conversion = subprocess.run(
                            [
                                soffice,
                                "--headless",
                                "--convert-to",
                                "txt:Text",
                                str(tmp_path),
                                "--outdir",
                                tmp_dir,
                            ],
                            capture_output=True,
                            text=True,
                        )

                    if conversion.returncode == 0:
                        out_path = Path(tmp_dir) / f"{tmp_path.stem}.txt"
                        if out_path.exists():
                            with open(out_path, "r", encoding="utf-8", errors="ignore") as f:
                                content = f.read().strip()

                            logger.info(
                                "DOC extraction completed via soffice",
                                extra_data={
                                    "file_name": file_name,
                                    "characters_extracted": len(content),
                                    "extraction_time_ms": timer.get_elapsed_ms(),
                                },
                            )
                            return content

            raise ExtractionError(
                "Failed to extract .doc file. Install textutil (macOS) or LibreOffice, or convert to DOCX."
            )
        finally:
            try:
                tmp_path.unlink()
            except Exception:
                pass

    def _extract_xls(self, file_bytes: bytes, file_name: str = "unknown.xls") -> str:
        """Extract text from legacy .xls using xlrd."""
        try:
            import xlrd
        except ImportError as exc:
            raise ExtractionError(
                "xlrd is required for .xls extraction. Please install xlrd."
            ) from exc

        try:
            with Timer("xls_extraction") as timer:
                workbook = xlrd.open_workbook(file_contents=file_bytes)
                parts: list[str] = []

                for sheet in workbook.sheets():
                    parts.append(f"## Sheet: {sheet.name}")
                    for row_idx in range(sheet.nrows):
                        row_values = []
                        for cell_value in sheet.row_values(row_idx):
                            if isinstance(cell_value, float) and cell_value.is_integer():
                                cell_str = str(int(cell_value))
                            else:
                                cell_str = str(cell_value)
                            row_values.append(cell_str.strip())
                        parts.append(" | ".join(row_values))

                result = "\n".join(parts).strip()

            logger.info(
                "XLS extraction completed",
                extra_data={
                    "file_name": file_name,
                    "sheet_count": len(workbook.sheets()),
                    "characters_extracted": len(result),
                    "extraction_time_ms": timer.get_elapsed_ms(),
                },
            )

            return result
        except Exception as exc:
            logger.error(
                "XLS extraction failed",
                extra_data={
                    "file_name": file_name,
                    "error_type": type(exc).__name__,
                    "error": str(exc),
                },
                exc_info=True,
            )
            raise ExtractionError(f"Failed to extract .xls file: {exc}") from exc

    def _extract_xlsx(self, file_bytes: bytes, file_name: str = "unknown.xlsx") -> str:
        """Extract text from modern .xlsx files using openpyxl."""
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise ExtractionError(
                "openpyxl is required for .xlsx extraction. Please install openpyxl."
            ) from exc

        try:
            with Timer("xlsx_extraction") as timer:
                # Load workbook from bytes
                workbook = load_workbook(io.BytesIO(file_bytes), data_only=True)
                parts: list[str] = []

                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    parts.append(f"## Sheet: {sheet_name}")

                    for row in sheet.iter_rows(values_only=True):
                        row_values = []
                        for cell_value in row:
                            if cell_value is None:
                                cell_str = ""
                            elif isinstance(cell_value, float) and cell_value.is_integer():
                                cell_str = str(int(cell_value))
                            else:
                                cell_str = str(cell_value)
                            row_values.append(cell_str.strip())
                        parts.append(" | ".join(row_values))

                result = "\n".join(parts).strip()

            logger.info(
                "XLSX extraction completed",
                extra_data={
                    "file_name": file_name,
                    "sheet_count": len(workbook.sheetnames),
                    "characters_extracted": len(result),
                    "extraction_time_ms": timer.get_elapsed_ms(),
                },
            )

            return result
        except Exception as exc:
            logger.error(
                "XLSX extraction failed",
                extra_data={
                    "file_name": file_name,
                    "error_type": type(exc).__name__,
                    "error": str(exc),
                },
                exc_info=True,
            )
            raise ExtractionError(f"Failed to extract .xlsx file: {exc}") from exc

    def _extract_image(self, file_bytes: bytes, file_name: str = "unknown.jpg") -> str:
        """Extract from image using Tesseract OCR directly."""
        try:
            # Open image with PIL
            image = Image.open(io.BytesIO(file_bytes))
            image_format = image.format
            image_size = image.size

            logger.debug(
                "Starting OCR on image",
                extra_data={
                    "file_name": file_name,
                    "image_format": image_format,
                    "image_width": image_size[0],
                    "image_height": image_size[1],
                },
            )

            # Perform OCR with configured languages
            with Timer("image_ocr") as timer:
                text = pytesseract.image_to_string(
                    image,
                    lang=self.config.languages,
                    config=f"--psm {self.config.psm_mode}",
                )

            result = text.strip()

            logger.info(
                "Image OCR completed",
                extra_data={
                    "file_name": file_name,
                    "image_format": image_format,
                    "image_dimensions": f"{image_size[0]}x{image_size[1]}",
                    "characters_extracted": len(result),
                    "ocr_time_ms": timer.get_elapsed_ms(),
                },
            )

            return result

        except Exception as e:
            logger.error(
                "Image OCR failed",
                extra_data={
                    "file_name": file_name,
                    "error_type": type(e).__name__,
                    "error": str(e),
                },
                exc_info=True,
            )
            raise ExtractionError(f"Failed to extract text from image: {e}")
